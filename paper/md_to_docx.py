"""Convert obsessive_coherence.md to formatted .docx for Zenodo publication."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

MD = Path(r"C:\Users\ohiod\Projects\kappa-sig\paper\obsessive_coherence.md")
OUT = Path(r"C:\Users\ohiod\Projects\kappa-sig\paper\Obsessive_Coherence_Ohio_2026.docx")

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(11)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

# Read markdown
text = MD.read_text(encoding='utf-8')
lines = text.split('\n')

def add_formatted_text(paragraph, text):
    """Add text with basic markdown formatting (bold, italic)."""
    # Process bold and italic
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

in_table = False
table_rows = []
in_blockquote = False

def flush_table():
    global table_rows, in_table
    if not table_rows:
        return
    # Filter separator rows
    data = [r for r in table_rows if not all(c.strip().replace('-','').replace('|','') == '' for c in r)]
    if not data:
        table_rows = []; in_table = False; return
    ncols = max(len(r) for r in data)
    t = doc.add_table(rows=len(data), cols=ncols)
    t.style = 'Light Shading'
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            if j < ncols:
                t.rows[i].cells[j].text = cell.strip()
                for p in t.rows[i].cells[j].paragraphs:
                    p.style.font.size = Pt(9)
    table_rows = []; in_table = False

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip horizontal rules
    if stripped == '---' or stripped == '***':
        if in_table: flush_table()
        i += 1; continue
    
    # Table rows
    if '|' in stripped and stripped.startswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        # Skip separator rows like |---|---|
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1; continue
        if not in_table:
            in_table = True
            table_rows = []
        table_rows.append(cells)
        i += 1; continue
    elif in_table:
        flush_table()

    # Title (# at top)
    if stripped.startswith('# ') and not stripped.startswith('## '):
        title = stripped[2:]
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue
    
    # Headings
    if stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
        i += 1; continue
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
        i += 1; continue
    
    # Blockquote (> formulas)
    if stripped.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(stripped[2:])
        run.italic = True
        run.font.size = Pt(10)
        i += 1; continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.+)', stripped)
    if m:
        p = doc.add_paragraph(style='List Number')
        add_formatted_text(p, m.group(2))
        i += 1; continue
    
    # Bullet list
    if stripped.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        add_formatted_text(p, stripped[2:])
        i += 1; continue
    
    # Empty line
    if not stripped:
        i += 1; continue
    
    # Regular paragraph
    p = doc.add_paragraph()
    add_formatted_text(p, stripped)
    i += 1

# Flush remaining table
if in_table: flush_table()

# Save
doc.save(str(OUT))
print(f"DOCX saved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
